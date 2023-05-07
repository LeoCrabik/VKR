from docxtpl import DocxTemplate
import docx
import os


class FillTemplate:
    def __init__(self, student_dict, global_data, exam_units, template_path='template'):
        self.template_path = template_path
        self.doc = DocxTemplate(os.path.join(self.template_path, 'template.docx'))
        self.student_dict = student_dict
        self.global_data = global_data
        self.exam_units = exam_units
        self.student_list = student_dict.keys()
        self.out_path = 'output'
        self.context = None
        self.qualification_dict = {
            'Бакалавриат': 'бакалавра',
            'Специалитет': 'специалиста',
            'Магистратура': 'магистра',
            'Аспирантура': 'аспиранта'}

    def prepare_data(self, name):
        self.context = self.global_data
        self.context['second_name'], self.context['first_name'], self.context['third_name'] = name.split(' ')
        self.context['birth_date'] = self.student_dict[name][2]
        self.context['qualification_with_mark'] = self.qualification_dict[self.global_data['qualification']] + self.student_dict[name][3]

        for value, subjects in enumerate(self.student_dict[name][1]):
            self.context['kp{}'.format(value)] = '{}:{}'.format(subjects[0], subjects[1])
            self.context['kpmark{}'.format(value)] = subjects[2]

        j = 0
        try:
            for i, subjects in enumerate(self.student_dict[name][0]):
                # print('' in self.exam_units.keys())
                if subjects[-1] == 0:
                    self.context['exam{}'.format(i)] = subjects[0]
                    if subjects[0] in self.exam_units.keys():
                        self.context['exam{}_unit'.format(i)] = '{} з.е.'.format(self.exam_units[subjects[0]]) if 'з.е' not in self.exam_units[subjects[0]] else self.exam_units[subjects[0]]
                        self.context['exam{}_mark'.format(i)] = subjects[1]
                    else:
                        self.context['exam{}_unit'.format(i)] = subjects[1]
                        self.context['exam{}_mark'.format(i)] = subjects[2]
                else:
                    self.context['exam_new{}'.format(j)] = subjects[0]
                    if subjects[0] in self.exam_units.keys():
                        self.context['exam_new{}_unit'.format(j)] = '{} з.е.'.format(self.exam_units[subjects[0]]) if 'з.е' not in self.exam_units[subjects[0]] else self.exam_units[subjects[0]]
                        self.context['exam_new{}_mark'.format(j)] = subjects[1]
                    else:
                        self.context['exam{}_unit'.format(j)] = subjects[1]
                        self.context['exam_new{}_mark'.format(j)] = subjects[2]
                    j += 1
        except Exception as e:
            raise e
        self.doc.render(context=self.context)
        self.save(name.split(' ')[0])

    def fill_words(self):
        try:
            os.mkdir(os.path.join(self.out_path, self.global_data['group']))
        except FileExistsError:
            pass

        for name in self.student_list:
            self.prepare_data(name)

    def save(self, filename):
        save_path = os.path.join(self.out_path, self.global_data['group'], '{}.docx'.format(filename))
        self.doc.save(save_path)
        doc = docx.Document(save_path)
        first_table = doc.tables[-1]._cells[0].tables[0]
        second_table = doc.tables[-1]._cells[2].tables[0]

        def remove_row(table, row):
            tbl = table._tbl
            tr = row._tr
            tbl.remove(tr)

        def delete_rows(table):
            while not table.rows[-1].cells[0].text:
                remove_row(table, table.rows[-1])

        delete_rows(first_table)
        delete_rows(second_table)

        doc.save(save_path)

